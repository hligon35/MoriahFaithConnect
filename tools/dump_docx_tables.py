from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python tools/dump_docx_tables.py <docx>")
        raise SystemExit(2)

    docx_path = Path(sys.argv[1])
    doc = Document(docx_path)

    print(f"tables: {len(doc.tables)}")
    for t_i, table in enumerate(doc.tables):
        print(f"\n=== Table {t_i} ===")
        for r_i, row in enumerate(table.rows):
            cells = [c.text.replace("\u00a0", " ").strip() for c in row.cells]
            # collapse empties for readability
            cells = [c for c in cells if c]
            if not cells:
                continue
            print(f"{r_i:03d}: " + " | ".join(cells))


if __name__ == "__main__":
    main()
