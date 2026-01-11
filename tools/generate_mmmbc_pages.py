from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup
from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = REPO_ROOT / "assets" / "webPages"
PAGES_DIR = REPO_ROOT / "mmmbc" / "Pages"
IMAGES_DIR = REPO_ROOT / "mmmbc" / "ConImg" / "webPages"


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _doc_paragraphs(docx_name: str) -> list[str]:
    doc = Document(DOCX_DIR / docx_name)
    paras = [_clean_text(p.text) for p in doc.paragraphs]
    return [p for p in paras if p]


def _replace_section_html(html_path: Path, new_section_inner_html: str) -> None:
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    section = soup.select_one("section.content-section")
    if section is None:
        raise RuntimeError(f"No .content-section found in {html_path}")

    section.clear()
    fragment = BeautifulSoup(new_section_inner_html, "html.parser")
    for child in list(fragment.contents):
        section.append(child)

    html_path.write_text(str(soup), encoding="utf-8")


def _img_tag(src: str, alt: str) -> str:
    return f'<img src="{src}" alt="{alt}">'


def build_church_history() -> str:
    paras = _doc_paragraphs("CHURCH HISTORY.docx")

    # Title usually first paragraph
    title = paras[0] if paras else "Church History"
    body = paras[1:]

    # Fix garbled first paragraph if present
    if body and body[0].lower().startswith("aand") and "Mt." in body[0]:
        body[0] = body[0].split("Mt.", 1)[-1]
        body[0] = "Mt." + body[0]

    imgs = [
        "../ConImg/webPages/church_history_image1.jpeg",
        "../ConImg/webPages/church_history_image2.jpeg",
        "../ConImg/webPages/church_history_image3.jpeg",
    ]

    out: list[str] = [f"<h1>{title}</h1>"]
    out.append('<div class="content-image-grid">')
    for src in imgs:
        out.append(_img_tag(src, "Mt. Moriah Missionary Baptist Church historical photo"))
    out.append("</div>")

    for p in body:
        out.append(f"<p>{p}</p>")

    return "\n".join(out)


@dataclass
class PersonSection:
    name: str
    paragraphs: list[str]


def _split_people_from_docx(docx_name: str) -> list[PersonSection]:
    paras = _doc_paragraphs(docx_name)

    # Identify header paragraphs (short lines starting with titles)
    header_prefixes = ("Evangelist", "Elder", "Rev.", "Reverend", "Minister", "Pastor")

    headers: list[tuple[int, str]] = []
    for idx, text in enumerate(paras):
        if text.startswith(header_prefixes) and len(text) <= 80:
            headers.append((idx, text))

    if not headers:
        return []

    sections: list[PersonSection] = []
    for h_i, h_text in headers:
        next_headers = [i for i, _ in headers if i > h_i]
        end = min(next_headers) if next_headers else len(paras)
        body = paras[h_i + 1 : end]
        # Split any embedded newlines that were collapsed into spaces but represent headings
        # We'll keep long items as paragraphs; allow manual sub-headings via markers.
        cleaned: list[str] = []
        for line in body:
            # If the original doc created "Education Bachelor..." style blocks, keep as-is.
            cleaned.append(line)
        sections.append(PersonSection(name=h_text, paragraphs=cleaned))

    return sections


def build_ministries() -> str:
    people = _split_people_from_docx("ministries.docx")

    image_files = [
        "ministries_image1.png",
        "ministries_image2.jpeg",
        "ministries_image3.png",
        "ministries_image4.jpeg",
        "ministries_image5.png",
        "ministries_image6.jpeg",
    ]

    out: list[str] = ["<h1>Ministries</h1>"]
    out.append(
        "<p>Learn more about the ministries and leaders who serve the Mt. Moriah Missionary Baptist Church family.</p>"
    )

    for i, person in enumerate(people):
        img = image_files[i] if i < len(image_files) else ""
        img_src = f"../ConImg/webPages/{img}" if img else ""

        out.append('<div class="leadership-profile">')
        if img_src:
            out.append(
                f'<img class="profile-image" src="{img_src}" alt="{person.name}">'  # noqa: S308
            )
        out.append('<div class="profile-details">')
        out.append(f"<h2>{person.name}</h2>")
        for p in person.paragraphs:
            out.append(f"<p>{p}</p>")
        out.append("</div>")
        out.append("</div>")

    return "\n".join(out)


def build_leadership() -> str:
    staff_paras = _doc_paragraphs("LEADERSHIP & STAFF.docx")
    deacons_paras = _doc_paragraphs("DEACONS.docx")
    deaconess_paras = _doc_paragraphs("DEACONESSES.docx")
    official_paras = _doc_paragraphs("OFFICIAL TEAM.docx")

    out: list[str] = ["<h1>Leadership &amp; Staff</h1>"]
    out.append(
        '<ul class="content-subnav">'
        '<li><a href="#staff">Staff</a></li>'
        '<li><a href="#deacons">Deacons</a></li>'
        '<li><a href="#deaconesses">Deaconesses</a></li>'
        '<li><a href="#official-team">Official Team &amp; Trustees</a></li>'
        "</ul>"
    )

    # Staff section: expect repeating pattern: Name, Title, Description
    out.append('<h2 id="staff">Staff</h2>')
    staff_imgs = [
        "leadership_staff_image1.jpeg",
        "leadership_staff_image2.png",
        "leadership_staff_image3.jpeg",
    ]

    # Parse as chunks: name line, title line, then description line(s) until next name
    def chunk_triplets(paras: list[str]) -> list[tuple[str, str, str]]:
        chunks: list[tuple[str, str, str]] = []
        i = 0
        while i < len(paras):
            name = paras[i]
            title = paras[i + 1] if i + 1 < len(paras) else ""
            desc = paras[i + 2] if i + 2 < len(paras) else ""
            chunks.append((name, title, desc))
            i += 3
        return chunks

    for i, (name, title, desc) in enumerate(chunk_triplets(staff_paras)):
        img = staff_imgs[i] if i < len(staff_imgs) else ""
        img_src = f"../ConImg/webPages/{img}" if img else ""
        out.append('<div class="leadership-profile">')
        if img_src:
            out.append(f'<img class="profile-image" src="{img_src}" alt="{name}">')
        out.append('<div class="profile-details">')
        out.append(f"<h2>{name}</h2>")
        if title:
            out.append(f"<p><strong>{title}</strong></p>")
        if desc:
            out.append(f"<p>{desc}</p>")
        out.append("</div>")
        out.append("</div>")

    # Deacons
    out.append('<h2 id="deacons">Deacons</h2>')
    if len(deacons_paras) >= 3:
        out.append(f"<p>{deacons_paras[1]}</p>")
        out.append(f"<p>{deacons_paras[2]}</p>")

    deacon_imgs = [
        "deacons_image1.jpeg",
        "deacons_image2.jpeg",
        "deacons_image3.jpeg",
        "deacons_image4.jpeg",
        "deacons_image5.png",
    ]

    # After intro, remaining are Name, Role, Bio repeating
    deacon_people: list[tuple[str, str, str]] = []
    rest = deacons_paras[3:]
    for i in range(0, len(rest), 3):
        part = rest[i : i + 3]
        if len(part) < 2:
            continue
        name = part[0]
        role = part[1] if len(part) > 1 else ""
        bio = part[2] if len(part) > 2 else ""
        deacon_people.append((name, role, bio))

    for i, (name, role, bio) in enumerate(deacon_people):
        img = deacon_imgs[i] if i < len(deacon_imgs) else ""
        img_src = f"../ConImg/webPages/{img}" if img else ""
        out.append('<div class="leadership-profile">')
        if img_src:
            out.append(f'<img class="profile-image" src="{img_src}" alt="{name}">')
        out.append('<div class="profile-details">')
        out.append(f"<h2>{name}</h2>")
        if role:
            out.append(f"<p><strong>{role}</strong></p>")
        if bio:
            out.append(f"<p>{bio}</p>")
        out.append("</div>")
        out.append("</div>")

    # Deaconesses
    out.append('<h2 id="deaconesses">Deaconesses</h2>')
    if len(deaconess_paras) >= 2:
        out.append(f"<p>{deaconess_paras[1]}</p>")

    deaconess_imgs = [
        "deaconesses_image1.jpeg",
        "deaconesses_image2.png",
        "deaconesses_image3.png",
        "deaconesses_image4.jpeg",
        "deaconesses_image5.png",
        "deaconesses_image6.png",
    ]

    # After intro: name, title, description repeating
    rest = deaconess_paras[2:]
    deaconess_people: list[tuple[str, str, str]] = []
    for i in range(0, len(rest), 3):
        part = rest[i : i + 3]
        if len(part) < 2:
            continue
        name = part[0]
        title = part[1] if len(part) > 1 else ""
        desc = part[2] if len(part) > 2 else ""
        deaconess_people.append((name, title, desc))

    for i, (name, title, desc) in enumerate(deaconess_people):
        img = deaconess_imgs[i] if i < len(deaconess_imgs) else ""
        img_src = f"../ConImg/webPages/{img}" if img else ""
        out.append('<div class="leadership-profile">')
        if img_src:
            out.append(f'<img class="profile-image" src="{img_src}" alt="{name}">')
        out.append('<div class="profile-details">')
        out.append(f"<h2>{name}</h2>")
        if title:
            out.append(f"<p><strong>{title}</strong></p>")
        if desc:
            out.append(f"<p>{desc}</p>")
        out.append("</div>")
        out.append("</div>")

    # Official team
    out.append('<h2 id="official-team">Official Team &amp; Trustees</h2>')
    if len(official_paras) >= 3:
        out.append(f"<p>{official_paras[1]}</p>")
        out.append(f"<p>{official_paras[2]}</p>")

    official_imgs = [
        "official_team_image1.png",
        "official_team_image2.png",
        "official_team_image3.png",
        "official_team_image4.png",
        "official_team_image5.png",
        "official_team_image6.png",
        "official_team_image7.jpeg",
    ]

    rest = official_paras[3:]
    official_people: list[tuple[str, str]] = []
    for i in range(0, len(rest), 2):
        part = rest[i : i + 2]
        if len(part) < 1:
            continue
        name = part[0]
        desc = part[1] if len(part) > 1 else ""
        official_people.append((name, desc))

    for i, (name, desc) in enumerate(official_people):
        img = official_imgs[i] if i < len(official_imgs) else ""
        img_src = f"../ConImg/webPages/{img}" if img else ""
        out.append('<div class="leadership-profile">')
        if img_src:
            out.append(f'<img class="profile-image" src="{img_src}" alt="{name}">')
        out.append('<div class="profile-details">')
        out.append(f"<h2>{name}</h2>")
        if desc:
            out.append(f"<p>{desc}</p>")
        out.append("</div>")
        out.append("</div>")

    return "\n".join(out)


def build_contact() -> str:
    paras = _doc_paragraphs("GET IN TOUCH.docx")

    # Extract key fields
    address = "1201 S 8th St, Paducah, KY 42003"
    phone = "(270) 443-3714"
    email = "mtmoriahmbc1201@gmail.com"
    fax = "(270) 443-7125"

    intro = "Have questions? Need more information about our church or an upcoming event? Please feel free to reach out to us!"
    if len(paras) >= 2:
        intro = paras[1].replace("Flease", "Please")

    out: list[str] = ["<h1>Contact Us</h1>"]
    out.append(f"<p>{intro}</p>")

    out.append("<h2>Get In Touch</h2>")
    out.append("<ul>")
    out.append(f'<li><strong>Visit Us:</strong> {address}</li>')
    out.append(f'<li><strong>Call Us:</strong> <a href="tel:2704433714">{phone}</a></li>')
    out.append(f'<li><strong>Email Us:</strong> <a href="mailto:{email}">{email}</a></li>')
    out.append(f'<li><strong>Fax:</strong> {fax}</li>')
    out.append("</ul>")

    out.append('<p>We value your connection. Please select an option below to send us a message or share a prayer request.</p>')

    # Keep the existing form markup by leaving placeholders that we will not overwrite here.
    # The generator will be used only to inject the new content ABOVE the forms.
    out.append("<!-- FORMS_PLACEHOLDER -->")

    return "\n".join(out)


def build_facility_rental() -> str:
    # The provided FACILITY RENTAL.docx appears to contain church history text.
    # We still surface a clean facility rental page using the extracted images.
    imgs = [
        "../ConImg/webPages/facility_rental_image1.jpeg",
        "../ConImg/webPages/facility_rental_image2.jpeg",
        "../ConImg/webPages/facility_rental_image3.jpeg",
    ]

    out: list[str] = ["<h1>Facility Rental</h1>"]
    out.append('<div class="content-image-grid">')
    for src in imgs:
        out.append(_img_tag(src, "Mt. Moriah facility"))
    out.append("</div>")

    out.append(
        "<p>Thank you for your interest in using Mt. Moriah Missionary Baptist Church facilities. Facility rental information and guidelines are currently being updated.</p>"
    )
    out.append(
        '<p>For availability, pricing, and requirements, please contact the church office via the <a href="contact.html">Contact Us</a> page or call (270) 443-3714.</p>'
    )

    return "\n".join(out)


def main() -> None:
    # Church History
    _replace_section_html(PAGES_DIR / "church_history.html", build_church_history())

    # Ministries
    _replace_section_html(PAGES_DIR / "ministries.html", build_ministries())

    # Leadership
    _replace_section_html(PAGES_DIR / "leadership.html", build_leadership())

    # Contact: preserve existing forms by splicing them back in
    contact_path = PAGES_DIR / "contact.html"
    soup = BeautifulSoup(contact_path.read_text(encoding="utf-8"), "html.parser")
    section = soup.select_one("section.content-section")
    if section is None:
        raise RuntimeError("No .content-section in contact.html")

    # capture existing toggle + forms
    existing_forms = section.select_one("div.form-toggle-buttons")
    prayer_form = section.select_one("form#prayerRequestForm")
    contact_form = section.select_one("form#contactInfoForm")

    new_html = build_contact()
    section.clear()
    fragment = BeautifulSoup(new_html, "html.parser")

    placeholder = None
    for node in fragment.contents:
        if getattr(node, "name", None) == "comment" and "FORMS_PLACEHOLDER" in str(node):
            placeholder = node
            break

    for child in list(fragment.contents):
        if child is placeholder:
            # insert the existing forms here
            if existing_forms is not None:
                section.append(existing_forms)
            if prayer_form is not None:
                section.append(prayer_form)
            if contact_form is not None:
                section.append(contact_form)
            continue
        section.append(child)

    contact_path.write_text(str(soup), encoding="utf-8")

    # Facility rental
    _replace_section_html(PAGES_DIR / "facility_rental.html", build_facility_rental())

    print("Updated pages: church_history, ministries, leadership, contact, facility_rental")


if __name__ == "__main__":
    main()
